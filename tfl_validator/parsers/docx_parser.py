"""Parse tables from Word (.docx) TFL files into structured DataFrames."""
import pandas as pd
from docx import Document


def extract_tables(filepath):
    """Extract all tables from a .docx file.
    Returns a list of DataFrames, one per table found.
    """
    doc = Document(filepath)
    tables = []

    for i, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)

        if len(rows) < 2:
            continue

        # First row as header
        header = rows[0]
        data = rows[1:]
        df = pd.DataFrame(data, columns=header)
        df.attrs["source_file"] = filepath
        df.attrs["table_index"] = i
        tables.append(df)

    return tables


def extract_all_text(filepath):
    """Extract all text from a .docx file (title, paragraphs, table cells)."""
    doc = Document(filepath)
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    return "\n".join(lines)


def extract_numbers_from_table(df):
    """Extract all numeric values from a table DataFrame.
    Returns dict mapping (row_label, col_header) -> value string.
    """
    numbers = {}
    for col in df.columns:
        for idx, row in df.iterrows():
            val = str(row[col]).strip()
            if val and val != "":
                # Use first column as row label if available
                row_label = str(row.iloc[0]) if len(df.columns) > 1 else str(idx)
                numbers[(row_label, col)] = val
    return numbers
